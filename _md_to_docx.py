# -*- coding: utf-8 -*-
"""Converte os .md do TCC e do relatorio de estagio em .docx (snapshot 'como esta')."""
import re, io, os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INLINE_RE = re.compile(r'(\*\*.+?\*\*|`.+?`|\[.+?\]\(.+?\)|\*.+?\*)')

def add_inline(paragraph, text):
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        tok = m.group(0)
        if tok.startswith('**') and tok.endswith('**'):
            r = paragraph.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith('`'):
            r = paragraph.add_run(tok[1:-1]); r.font.name = 'Consolas'; r.font.size = Pt(10)
        elif tok.startswith('['):
            mm = re.match(r'\[(.+?)\]\((.+?)\)', tok)
            paragraph.add_run(mm.group(1) if mm else tok)
        elif tok.startswith('*'):
            r = paragraph.add_run(tok[1:-1]); r.italic = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])

def is_table_sep(line):
    s = line.strip()
    return s.startswith('|') and bool(re.match(r'^[\s:\-|]+$', s)) and '-' in s

def add_hr(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    for k, v in (('w:val', 'single'), ('w:sz', '6'), ('w:space', '1'), ('w:color', '999999')):
        bottom.set(qn(k), v)
    pbdr.append(bottom); pPr.append(pbdr)

def add_image_box(doc, descricao):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    for edge in ('top', 'bottom', 'left', 'right'):
        e = OxmlElement('w:' + edge)
        for k, v in (('w:val', 'single'), ('w:sz', '6'), ('w:space', '6'), ('w:color', '999999')):
            e.set(qn(k), v)
        pbdr.append(e)
    pPr.append(pbdr)
    r = p.add_run('[ IMAGEM ]  ' + descricao)
    r.italic = True; r.font.color.rgb = RGBColor(0x66, 0x66, 0x66); r.font.size = Pt(11)

def cells_of(row):
    return [c.strip() for c in row.strip().strip('|').split('|')]

def parse_md(md_lines, doc):
    i, n = 0, len(md_lines)
    while i < n:
        raw = md_lines[i].rstrip('\n')
        s = raw.strip()
        if not s:
            i += 1; continue
        # caixa de imagem [[IMG: ...]]
        mimg = re.match(r'^\[\[IMG:\s*(.*?)\]\]$', s)
        if mimg:
            add_image_box(doc, mimg.group(1).strip()); i += 1; continue
        # code fence
        if s.startswith('```'):
            i += 1; code = []
            while i < n and not md_lines[i].strip().startswith('```'):
                code.append(md_lines[i].rstrip('\n')); i += 1
            i += 1
            p = doc.add_paragraph()
            r = p.add_run('\n'.join(code)); r.font.name = 'Consolas'; r.font.size = Pt(9)
            continue
        # heading
        m = re.match(r'^(#{1,6})\s+(.*)$', s)
        if m:
            doc.add_heading(m.group(2).strip(), level=min(len(m.group(1)), 4))
            i += 1; continue
        # horizontal rule
        if re.match(r'^(\*\*\*+|---+|___+)$', s):
            add_hr(doc); i += 1; continue
        # table
        if s.startswith('|') and i + 1 < n and is_table_sep(md_lines[i + 1]):
            rows = []
            while i < n and md_lines[i].strip().startswith('|'):
                rows.append(md_lines[i]); i += 1
            header = cells_of(rows[0]); ncol = len(header)
            data = [cells_of(r) for r in rows[2:]]
            t = doc.add_table(rows=1, cols=ncol); t.style = 'Table Grid'
            for c, htxt in enumerate(header):
                cell = t.rows[0].cells[c]
                add_inline(cell.paragraphs[0], htxt)
                for run in cell.paragraphs[0].runs:
                    run.bold = True
            for drow in data:
                cells = t.add_row().cells
                for c in range(ncol):
                    add_inline(cells[c].paragraphs[0], drow[c] if c < len(drow) else '')
            doc.add_paragraph()
            continue
        # blockquote
        if s.startswith('>'):
            q = []
            while i < n and md_lines[i].strip().startswith('>'):
                q.append(re.sub(r'^\s*>\s?', '', md_lines[i].rstrip('\n'))); i += 1
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Pt(18)
            add_inline(p, ' '.join(x for x in q if x.strip()))
            for run in p.runs:
                run.italic = True; run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            continue
        # bullet list
        m = re.match(r'^\s*[-*+]\s+(.*)$', raw)
        if m:
            add_inline(doc.add_paragraph(style='List Bullet'), m.group(1)); i += 1; continue
        # numbered list
        m = re.match(r'^\s*\d+\.\s+(.*)$', raw)
        if m:
            add_inline(doc.add_paragraph(style='List Number'), m.group(1)); i += 1; continue
        # paragraph
        para = [raw]; i += 1
        while i < n:
            nx = md_lines[i]; nxs = nx.strip()
            if (not nxs or nxs.startswith('#') or nxs.startswith('>') or nxs.startswith('|')
                    or nxs.startswith('```') or re.match(r'^\s*[-*+]\s', nx)
                    or re.match(r'^\s*\d+\.\s', nx) or re.match(r'^(\*\*\*+|---+|___+)$', nxs)):
                break
            para.append(nx.rstrip('\n')); i += 1
        add_inline(doc.add_paragraph(), ' '.join(x.strip() for x in para))

def add_footer_pagenum(doc):
    try:
        sec = doc.sections[0]
        p = sec.footer.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        f1 = OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'), 'begin')
        it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = 'PAGE'
        f2 = OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'), 'end')
        run._r.append(f1); run._r.append(it); run._r.append(f2)
    except Exception as e:
        print('  (footer skip:', e, ')')

def build(files, out_path, draft_note):
    doc = Document()
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(12)
    for lvl in range(1, 5):
        try:
            st = doc.styles[f'Heading {lvl}']
            st.font.color.rgb = RGBColor(0, 0, 0); st.font.name = 'Arial'
        except KeyError:
            pass
    first = True
    for f in files:
        if not os.path.exists(f):
            print('  MISSING:', f); continue
        if not first:
            doc.add_page_break()
        first = False
        with io.open(f, encoding='utf-8') as fh:
            parse_md(fh.read().split('\n'), doc)
    add_footer_pagenum(doc)
    doc.save(out_path)
    d2 = Document(out_path)
    print('  OK ->', out_path, '| paragrafos:', len(d2.paragraphs), '| tabelas:', len(d2.tables))

TCC = [
    'tcc/00_capa.md', 'tcc/04_resumo.md',
    'tcc/capitulos/1_introducao.md', 'tcc/capitulos/2_revisao_literatura.md',
    'tcc/capitulos/3_metodologia.md',
    # Resultados, Discussao e Conclusao OMITIDOS de proposito (TCC em andamento)
    'tcc/07_referencias.md', 'tcc/08_perguntas.md',
]
REL = [
    'relatorio_estagio/00_capa.md', 'relatorio_estagio/01_identificacao.md',
    'relatorio_estagio/secoes/1_introducao.md', 'relatorio_estagio/secoes/2_empresa.md',
    'relatorio_estagio/secoes/3_atividades.md', 'relatorio_estagio/secoes/4_resultados.md',
    'relatorio_estagio/secoes/5_consideracoes.md',
]

print('TCC:')
build(TCC, 'tcc/TCC_PCP_Vo_Nena_2026-06-01.docx', 'TCC')
print('FEITO. (o relatorio e gerado em ABNT por _gerar_relatorio_abnt.py)')
