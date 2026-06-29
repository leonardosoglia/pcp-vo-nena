# -*- coding: utf-8 -*-
"""Gera o TCC completo em formato ABNT (NBR 14724) a partir dos .md.
Mesma engine ABNT do _gerar_relatorio_abnt.py, montagem propria do TCC:
capa, folha de rosto, resumo, abstract, sumario, caps 1-6 e referencias.
NAO sobrescreve o .docx antigo (nome com a data de hoje)."""
import re, io, os
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TNR = 'Arial'  # fonte do documento (modelo do professor + ABNT usam Arial)
INLINE_RE = re.compile(r'(\*\*.+?\*\*|`.+?`|\[.+?\]\(.+?\)|\*.+?\*)')
SKIP_HEAD = ('notas', 'ainda falta', 'alternativas de t')

TITULO = ('DESENVOLVIMENTO DE UM SISTEMA DIGITAL DE PLANEJAMENTO E CONTROLE DA '
          'PRODUÇÃO PARA UMA CONFEITARIA INDUSTRIAL: DO PAPEL AO ALGORITMO COGNITIVO')


def setup_page(sec):
    sec.page_width, sec.page_height = Cm(21), Cm(29.7)
    sec.top_margin, sec.left_margin = Cm(3), Cm(3)
    sec.bottom_margin, sec.right_margin = Cm(2), Cm(2)
    sec.header_distance, sec.footer_distance = Cm(2), Cm(2)


def setup_styles(doc):
    nm = doc.styles['Normal']
    nm.font.name = TNR; nm.font.size = Pt(12)
    pf = nm.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.first_line_indent = Cm(1.25)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_before = Pt(0); pf.space_after = Pt(0)
    for lvl in range(1, 5):
        st = doc.styles[f'Heading {lvl}']
        st.font.name = TNR; st.font.size = Pt(12); st.font.bold = True
        st.font.color.rgb = RGBColor(0, 0, 0); st.font.italic = (lvl >= 3)
        h = st.paragraph_format
        h.first_line_indent = Cm(0); h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        h.space_before = Pt(18); h.space_after = Pt(12)
        h.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        h.keep_with_next = True
    # estilo proprio das legendas de figura (coletado pela Lista de Figuras)
    try:
        leg = doc.styles['LegendaFigura']
    except KeyError:
        leg = doc.styles.add_style('LegendaFigura', WD_STYLE_TYPE.PARAGRAPH)
    leg.base_style = doc.styles['Normal']
    leg.font.name = TNR; leg.font.size = Pt(11); leg.font.bold = True
    lp = leg.paragraph_format
    lp.first_line_indent = Cm(0); lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lp.space_before = Pt(6); lp.space_after = Pt(2)
    lp.line_spacing_rule = WD_LINE_SPACING.SINGLE
    lp.keep_with_next = True
    # estilo das legendas de quadro (coletado pela Lista de Quadros)
    try:
        lq = doc.styles['LegendaQuadro']
    except KeyError:
        lq = doc.styles.add_style('LegendaQuadro', WD_STYLE_TYPE.PARAGRAPH)
    lq.base_style = doc.styles['Normal']
    lq.font.name = TNR; lq.font.size = Pt(11); lq.font.bold = True
    qp = lq.paragraph_format
    qp.first_line_indent = Cm(0); qp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    qp.space_before = Pt(6); qp.space_after = Pt(2)
    qp.line_spacing_rule = WD_LINE_SPACING.SINGLE
    qp.keep_with_next = True


def _field(run, instr):
    for typ, txt in (('begin', None), (None, instr), ('end', None)):
        if typ:
            fc = OxmlElement('w:fldChar'); fc.set(qn('w:fldCharType'), typ); run._r.append(fc)
        else:
            it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = txt
            run._r.append(it)


def add_pagenum(sec):
    sec.header.is_linked_to_previous = False
    p = sec.header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run(); r.font.name = TNR; r.font.size = Pt(10)
    _field(r, 'PAGE')


def centered(doc, text, bold=False, size=12, before=0, after=0):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.first_line_indent = Cm(0); pf.space_before = Pt(before); pf.space_after = Pt(after)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r = p.add_run(text); r.bold = bold; r.font.name = TNR; r.font.size = Pt(size)
    return p


def blanks(doc, n):
    for _ in range(n):
        centered(doc, '')


def add_image_box(doc, descricao):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    for edge in ('top', 'bottom', 'left', 'right'):
        e = OxmlElement('w:' + edge)
        for k, v in (('w:val', 'single'), ('w:sz', '6'), ('w:space', '6'), ('w:color', '999999')):
            e.set(qn(k), v)
        pbdr.append(e)
    pPr.append(pbdr)
    r = p.add_run('[ IMAGEM ]  ' + descricao)
    r.italic = True; r.font.color.rgb = RGBColor(0x66, 0x66, 0x66); r.font.size = Pt(11)


def add_inline(p, text):
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            p.add_run(text[pos:m.start()])
        tok = m.group(0)
        if tok.startswith('**') and tok.endswith('**'):
            p.add_run(tok[2:-2]).bold = True
        elif tok.startswith('`'):
            r = p.add_run(tok[1:-1]); r.font.name = 'Consolas'; r.font.size = Pt(10)
        elif tok.startswith('['):
            mm = re.match(r'\[(.+?)\]\((.+?)\)', tok); p.add_run(mm.group(1) if mm else tok)
        elif tok.startswith('*'):
            p.add_run(tok[1:-1]).italic = True
        pos = m.end()
    if pos < len(text):
        p.add_run(text[pos:])


def is_sep(line):
    s = line.strip()
    return s.startswith('|') and bool(re.match(r'^[\s:\-|]+$', s)) and '-' in s


def cells_of(row):
    return [c.strip() for c in row.strip().strip('|').split('|')]


def parse_md(lines, doc, clean=False, head_as_bold=False):
    i, n, skip = 0, len(lines), False
    while i < n:
        raw = lines[i].rstrip('\n'); s = raw.strip()
        mh = re.match(r'^(#{1,6})\s+(.*)$', s)
        if mh:
            title = mh.group(2).strip(); lvl = len(mh.group(1))
            if clean and any(title.lower().startswith(p) for p in SKIP_HEAD):
                skip = True; i += 1; continue
            skip = False
            if head_as_bold:
                p = doc.add_paragraph()
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.space_before = Pt(12 if lvl == 1 else 6)
                if lvl == 1:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r = p.add_run(title.upper())
                else:
                    r = p.add_run(title)
                r.bold = True; r.font.name = TNR; r.font.size = Pt(12)
            else:
                h = doc.add_heading('', level=min(lvl, 4))
                add_inline(h, title)
            i += 1; continue
        if skip:
            i += 1; continue
        if not s:
            i += 1; continue
        mimg = re.match(r'^\[\[IMG:\s*(.*?)\]\]$', s)
        if mimg:
            add_image_box(doc, mimg.group(1).strip()); i += 1; continue
        mfig = re.match(r'^\*\*\s*(Figura\s+\d+\s*[—–-]\s*.*?)\s*\*\*$', s)
        if mfig:
            p = doc.add_paragraph(style='LegendaFigura')
            r = p.add_run(mfig.group(1).strip())
            r.bold = True; r.font.name = TNR; r.font.size = Pt(11)
            i += 1; continue
        mquad = re.match(r'^\*\*\s*(Quadro\s+\d+\s*[—–-]\s*.*?)\s*\*\*$', s)
        if mquad:
            p = doc.add_paragraph(style='LegendaQuadro')
            r = p.add_run(mquad.group(1).strip())
            r.bold = True; r.font.name = TNR; r.font.size = Pt(11)
            i += 1; continue
        mfonte = re.match(r'^Fonte:\s', s)
        if mfonte:
            p = doc.add_paragraph(); pf = p.paragraph_format
            pf.first_line_indent = Cm(0); pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf.space_before = Pt(0); pf.space_after = Pt(6)
            pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
            r = p.add_run(s); r.font.name = TNR; r.font.size = Pt(10)
            i += 1; continue
        if s.startswith('```'):
            i += 1; code = []
            while i < n and not lines[i].strip().startswith('```'):
                code.append(lines[i].rstrip('\n')); i += 1
            i += 1
            p = doc.add_paragraph(); p.paragraph_format.first_line_indent = Cm(0)
            r = p.add_run('\n'.join(code)); r.font.name = 'Consolas'; r.font.size = Pt(9)
            continue
        if re.match(r'^(\*\*\*+|---+|___+)$', s):
            i += 1; continue
        if s.startswith('|') and i + 1 < n and is_sep(lines[i + 1]):
            rows = []
            while i < n and lines[i].strip().startswith('|'):
                rows.append(lines[i]); i += 1
            header = cells_of(rows[0]); ncol = len(header)
            data = [cells_of(r) for r in rows[2:]]
            t = doc.add_table(rows=1, cols=ncol); t.style = 'Table Grid'
            for c, h in enumerate(header):
                cell = t.rows[0].cells[c]; cell.paragraphs[0].paragraph_format.first_line_indent = Cm(0)
                add_inline(cell.paragraphs[0], h)
                for rn in cell.paragraphs[0].runs:
                    rn.bold = True; rn.font.size = Pt(11)
            for drow in data:
                cs = t.add_row().cells
                for c in range(ncol):
                    cs[c].paragraphs[0].paragraph_format.first_line_indent = Cm(0)
                    add_inline(cs[c].paragraphs[0], drow[c] if c < len(drow) else '')
                    for rn in cs[c].paragraphs[0].runs:
                        rn.font.size = Pt(11)
            doc.add_paragraph()
            continue
        if s.startswith('>'):
            if clean:
                while i < n and lines[i].strip().startswith('>'):
                    i += 1
                continue
            q = []
            while i < n and lines[i].strip().startswith('>'):
                q.append(re.sub(r'^\s*>\s?', '', lines[i].rstrip('\n'))); i += 1
            add_inline(doc.add_paragraph(), ' '.join(x for x in q if x.strip())); continue
        m = re.match(r'^\s*[-*+]\s+(.*)$', raw)
        if m:
            add_inline(doc.add_paragraph(style='List Bullet'), m.group(1)); i += 1; continue
        m = re.match(r'^\s*\d+\.\s+(.*)$', raw)
        if m:
            add_inline(doc.add_paragraph(style='List Number'), m.group(1)); i += 1; continue
        para = [raw]; i += 1
        while i < n:
            nx = lines[i]; nxs = nx.strip()
            if (not nxs or nxs.startswith('#') or nxs.startswith('>') or nxs.startswith('|')
                    or nxs.startswith('```') or re.match(r'^\s*[-*+]\s', nx)
                    or re.match(r'^\s*\d+\.\s', nx) or re.match(r'^(\*\*\*+|---+|___+)$', nxs)):
                break
            para.append(nx.rstrip('\n')); i += 1
        add_inline(doc.add_paragraph(), ' '.join(x.strip() for x in para))


def read(f):
    return io.open(f, encoding='utf-8').read().split('\n') if os.path.exists(f) else ['<<arquivo ausente: %s>>' % f]


def toc_field(doc, instr, placeholder):
    """Campo de indice (TOC) cujo placeholder e o RESULTADO do campo — ou seja,
    some quando o usuario atualiza o campo no Word (em vez de sobrar como texto solto)."""
    p = doc.add_paragraph(); p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run()
    fb = OxmlElement('w:fldChar'); fb.set(qn('w:fldCharType'), 'begin'); r._r.append(fb)
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = instr
    r._r.append(it)
    fs = OxmlElement('w:fldChar'); fs.set(qn('w:fldCharType'), 'separate'); r._r.append(fs)
    rt = p.add_run(placeholder); rt.italic = True; rt.font.size = Pt(11)
    r2 = p.add_run()
    fe = OxmlElement('w:fldChar'); fe.set(qn('w:fldCharType'), 'end'); r2._r.append(fe)
    return p


# ---------- montagem do TCC ----------
doc = Document()
setup_styles(doc)
setup_page(doc.sections[0])

# CAPA
logo = doc.add_paragraph(); logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
logo.paragraph_format.first_line_indent = Cm(0); logo.paragraph_format.space_after = Pt(0)
if os.path.exists('assets/ufcg_brasao.png'):
    logo.add_run().add_picture('assets/ufcg_brasao.png', width=Cm(2.6))
centered(doc, 'UNIVERSIDADE FEDERAL DE CAMPINA GRANDE - UFCG')
centered(doc, 'CENTRO DE CIÊNCIAS E TECNOLOGIA - CCT')
centered(doc, 'UNIDADE ACADÊMICA DE ENGENHARIA DE PRODUÇÃO - UAEP')
centered(doc, 'CURSO DE GRADUAÇÃO EM ENGENHARIA DE PRODUÇÃO')
blanks(doc, 9)
centered(doc, 'LEONARDO SÓGLIA')
blanks(doc, 8)
centered(doc, TITULO, size=12)
blanks(doc, 14)
centered(doc, 'Campina Grande – PB')
centered(doc, '2026')

# FOLHA DE ROSTO
doc.add_page_break()
centered(doc, 'LEONARDO SÓGLIA')
blanks(doc, 7)
centered(doc, TITULO, size=12)
blanks(doc, 6)
nat = doc.add_paragraph()
nat.paragraph_format.left_indent = Cm(8); nat.paragraph_format.first_line_indent = Cm(0)
nat.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
nat.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
rn = nat.add_run('Trabalho de Conclusão de Curso apresentado ao Curso de Graduação em '
                 'Engenharia de Produção da Universidade Federal de Campina Grande, como '
                 'requisito parcial para a obtenção do título de Bacharel em Engenharia '
                 'de Produção.')
rn.font.name = TNR; rn.font.size = Pt(10)
blanks(doc, 3)
for line in ('Orientador: Prof. Dr. Francisco Kegenaldo Alves de Sousa',):
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(8)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r = p.add_run(line); r.font.name = TNR; r.font.size = Pt(12)
blanks(doc, 13)
centered(doc, 'Campina Grande – PB')
centered(doc, '2026')

# FOLHA DE APROVACAO
doc.add_page_break()
centered(doc, 'LEONARDO SÓGLIA')
blanks(doc, 5)
centered(doc, TITULO, size=12)
blanks(doc, 5)
ap = doc.add_paragraph()
ap.paragraph_format.left_indent = Cm(8); ap.paragraph_format.first_line_indent = Cm(0)
ap.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
ap.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
rap = ap.add_run('Trabalho de Conclusão de Curso apresentado ao Curso de Graduação em '
                 'Engenharia de Produção da Universidade Federal de Campina Grande, como '
                 'requisito parcial para a obtenção do título de Bacharel em Engenharia '
                 'de Produção.')
rap.font.name = TNR; rap.font.size = Pt(10)
blanks(doc, 4)
centered(doc, 'Campina Grande, ______ de ____________________ de 2026.')
blanks(doc, 5)
centered(doc, 'BANCA EXAMINADORA')
blanks(doc, 4)
for nome, papel in (
        ('Prof. Dr. Francisco Kegenaldo Alves de Sousa', 'Orientador – UAEP/CCT/UFCG'),
        ('[Nome do(a) avaliador(a)]', 'Avaliador(a) – UAEP/CCT/UFCG'),
        ('[Nome do(a) avaliador(a)]', 'Avaliador(a) – UAEP/CCT/UFCG')):
    centered(doc, '__________________________________________________')
    centered(doc, nome)
    centered(doc, papel)
    blanks(doc, 2)

# RESUMO + ABSTRACT (pre-textuais)
doc.add_page_break()
parse_md(read('tcc/04_resumo.md'), doc, clean=True, head_as_bold=True)
doc.add_page_break()
parse_md(read('tcc/05_abstract.md'), doc, clean=True, head_as_bold=True)

# LISTA DE FIGURAS (campo Sumario de Ilustracoes — coleta o estilo LegendaFigura)
doc.add_page_break()
centered(doc, 'LISTA DE FIGURAS', bold=True, size=12, after=12)
toc_field(doc, 'TOC \\h \\z \\t "LegendaFigura;1"',
          '[Abra no Word > Atualizar campo para gerar a lista de figuras, depois de inserir as imagens.]')

# LISTA DE QUADROS (campo — coleta o estilo LegendaQuadro)
doc.add_page_break()
centered(doc, 'LISTA DE QUADROS', bold=True, size=12, after=12)
toc_field(doc, 'TOC \\h \\z \\t "LegendaQuadro;1"',
          '[Abra no Word > Atualizar campo para gerar a lista de quadros.]')

# LISTA DE ABREVIATURAS E SIGLAS
doc.add_page_break()
centered(doc, 'LISTA DE ABREVIATURAS E SIGLAS', bold=True, size=12, after=12)
SIGLAS = [
    ('ABC', 'Análise ABC (curva de Pareto)'),
    ('BOM', 'Bill of Materials (lista de materiais)'),
    ('CCT', 'Centro de Ciências e Tecnologia'),
    ('ERP', 'Enterprise Resource Planning (sistema integrado de gestão empresarial)'),
    ('IA', 'Inteligência Artificial'),
    ('LLM', 'Large Language Model (modelo de linguagem de grande porte)'),
    ('MAPE', 'Mean Absolute Percentage Error (erro percentual absoluto médio)'),
    ('MRP', 'Material Requirements Planning (planejamento das necessidades de materiais)'),
    ('NF-e', 'Nota Fiscal eletrônica'),
    ('OP', 'Ordem de Produção'),
    ('PCP', 'Planejamento e Controle da Produção'),
    ('PDV', 'Ponto de Venda'),
    ('PMI', 'Pequenas e Médias Indústrias'),
    ('TCC', 'Trabalho de Conclusão de Curso'),
    ('UAEP', 'Unidade Acadêmica de Engenharia de Produção'),
    ('UFCG', 'Universidade Federal de Campina Grande'),
    ('XML', 'eXtensible Markup Language'),
]
tsig = doc.add_table(rows=0, cols=2)
for sig, sig_def in SIGLAS:
    cells = tsig.add_row().cells
    cells[0].width = Cm(3); cells[1].width = Cm(13)
    pp0 = cells[0].paragraphs[0]; pp0.paragraph_format.first_line_indent = Cm(0)
    pp0.paragraph_format.space_after = Pt(2)
    r0 = pp0.add_run(sig); r0.font.name = TNR; r0.font.size = Pt(12)
    pp1 = cells[1].paragraphs[0]; pp1.paragraph_format.first_line_indent = Cm(0)
    pp1.paragraph_format.space_after = Pt(2)
    r1 = pp1.add_run(sig_def); r1.font.name = TNR; r1.font.size = Pt(12)

# SUMARIO
doc.add_page_break()
centered(doc, 'SUMÁRIO', bold=True, size=12, after=12)
toc_field(doc, 'TOC \\o "1-3" \\h \\z \\u',
          '[Abra no Word > Atualizar campo para gerar o sumário.]')

# ---- secao textual (com numeracao de pagina) ----
sec2 = doc.add_section(WD_SECTION.NEW_PAGE)
setup_page(sec2)
add_pagenum(sec2)

BODY = ['tcc/capitulos/1_introducao.md',
        'tcc/capitulos/2_revisao_literatura.md',
        'tcc/capitulos/3_metodologia.md',
        'tcc/capitulos/4_resultados.md',
        'tcc/capitulos/5_discussao.md',
        'tcc/capitulos/6_conclusao.md',
        'tcc/07_referencias.md']
for idx, f in enumerate(BODY):
    if idx > 0:
        doc.add_page_break()
    parse_md(read(f), doc, clean=True)

out = 'tcc/TCC_PCP_Vo_Nena_2026-06-29.docx'
doc.save(out)
chk = Document(out)
print('OK ->', out)
print('  secoes:', len(chk.sections), '| paragrafos:', len(chk.paragraphs), '| tabelas:', len(chk.tables))
hs = [pp.text for pp in chk.paragraphs if pp.style.name.startswith('Heading')]
print('  titulos numerados:', len(hs))
for h in hs[:20]:
    print('   -', h)
figs = [pp.text for pp in chk.paragraphs if pp.style.name == 'LegendaFigura']
print('  legendas de figura:', len(figs))
for fpar in figs:
    print('   *', fpar)
